from pathlib import Path
from docx import Document
from docx.table import Table
from lxml import etree

from .utils import get_img_dir
from ..formulas import OMML_NS, FormulaRef, ExtractionResult

_OMATH_PARA = f"{{{OMML_NS}}}oMathPara"
_OMATH = f"{{{OMML_NS}}}oMath"


def extract_docx(file_path: Path, temp_dir: Path) -> ExtractionResult:
    doc = Document(str(file_path))
    img_dir = get_img_dir(temp_dir, file_path)

    # Extract images
    img_paths = []
    for i, rel in enumerate(doc.part.rels.values()):
        if "image" in rel.reltype and not rel.is_external:
            ext = Path(rel.target_ref).suffix
            img_path = img_dir / f"img_{i}{ext}"
            img_path.write_bytes(rel.target_part.blob)
            img_paths.append(img_path)

    lines = []
    img_idx = 0
    formulas: list[FormulaRef] = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]
        if tag == "p":
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            if para is None:
                continue

            # Check for inline images
            if para._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing") or \
               para._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict"):
                if img_idx < len(img_paths):
                    lines.append(f"\n![image]({img_paths[img_idx].resolve()})\n")
                    img_idx += 1

            # Detect formulas in this paragraph
            para_formulas = _detect_para_formulas(para, formulas, len(formulas))
            if para_formulas:
                # Build paragraph text with formula markers
                text = _build_text_with_markers(para, para_formulas, len(formulas) - len(para_formulas))
            else:
                text = para.text.strip()

            if not text:
                continue

            style = para.style.name if para.style else ""
            if "Heading 1" in style:
                lines.append(f"\n# {text}\n")
            elif "Heading 2" in style:
                lines.append(f"\n## {text}\n")
            elif "Heading 3" in style:
                lines.append(f"\n### {text}\n")
            else:
                lines.append(text)

        elif tag == "tbl":
            for table in doc.tables:
                if table._tbl is element:
                    lines.append(_table_to_md(table))
                    break

    # Append any remaining images not matched inline
    while img_idx < len(img_paths):
        lines.append(f"\n![image]({img_paths[img_idx].resolve()})\n")
        img_idx += 1

    return ExtractionResult(
        markdown="\n".join(lines),
        formulas=formulas,
    )


def _detect_para_formulas(para, formulas_list: list, start_idx: int) -> list[FormulaRef]:
    """Detect OMML formulas in a paragraph and add them to the formulas list."""
    elem = para._element
    found = []

    # Check for display-mode formulas: <m:oMathPara>
    omath_paras = elem.findall(f".//{_OMATH_PARA}")
    for omp in omath_paras:
        xml_str = etree.tostring(omp, encoding="unicode")
        ref = FormulaRef(
            omml_xml=xml_str,
            is_inline=False,
            context_text=para.text.strip()[:200],
        )
        formulas_list.append(ref)
        found.append(ref)

    # Check for inline formulas: bare <m:oMath> not inside <m:oMathPara>
    all_omath = elem.findall(f".//{_OMATH}")
    for om in all_omath:
        parent = om.getparent()
        if parent is not None and parent.tag == _OMATH_PARA:
            continue
        xml_str = etree.tostring(om, encoding="unicode")
        ref = FormulaRef(
            omml_xml=xml_str,
            is_inline=True,
            context_text=para.text.strip()[:200],
        )
        formulas_list.append(ref)
        found.append(ref)

    return found


def _build_text_with_markers(para, para_formulas: list[FormulaRef], start_idx: int) -> str:
    """Build paragraph text replacing formula positions with markers.

    For display formulas (oMathPara), the entire paragraph content is the formula.
    For inline formulas, we insert the marker in place of the formula text.
    """
    # If there's a display formula, the whole paragraph is the formula
    for i, ref in enumerate(para_formulas):
        if not ref.is_inline:
            return ref.marker(start_idx + i)

    # For inline formulas, construct text with markers
    # Simple approach: get paragraph text and append markers for each inline formula
    # since python-docx text property strips math, we build from runs
    text_parts = []
    formula_idx = 0
    elem = para._element

    # Walk child elements in order
    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for child in elem:
        tag = child.tag.split("}")[-1]
        if tag == "r":
            # Regular run - extract text
            for t in child.findall(f"{{{ns_w}}}t"):
                if t.text:
                    text_parts.append(t.text)
        elif child.tag == _OMATH and formula_idx < len(para_formulas):
            # Inline formula - insert marker
            text_parts.append(f" {para_formulas[formula_idx].marker(start_idx + formula_idx)} ")
            formula_idx += 1
        elif child.tag == _OMATH_PARA and formula_idx < len(para_formulas):
            text_parts.append(f" {para_formulas[formula_idx].marker(start_idx + formula_idx)} ")
            formula_idx += 1

    result = "".join(text_parts).strip()
    return result if result else para.text.strip()


def _table_to_md(table: Table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)
    if not rows:
        return ""
    # Deduplicate merged cells (python-docx repeats them)
    clean_rows = []
    for row in rows:
        seen = []
        prev = None
        for c in row:
            if c == prev:
                seen.append("")
            else:
                seen.append(c)
            prev = c
        clean_rows.append(seen)
    header = "| " + " | ".join(clean_rows[0]) + " |"
    sep = "| " + " | ".join("---" for _ in clean_rows[0]) + " |"
    body = "\n".join("| " + " | ".join(r) + " |" for r in clean_rows[1:])
    return f"\n{header}\n{sep}\n{body}\n"
