"""Detect OMML math formulas in DOCX documents."""

import logging
from lxml import etree
from docx import Document
from pathlib import Path

from . import OMML_NS, FormulaRef

log = logging.getLogger("knowledge_extractor")

# XPath patterns for OMML elements
_OMATH_PARA = f"{{{OMML_NS}}}oMathPara"
_OMATH = f"{{{OMML_NS}}}oMath"


def detect_formulas(doc: Document) -> list[FormulaRef]:
    """Walk a DOCX document and detect all OMML formula elements.

    Returns a list of FormulaRef objects with OMML XML, inline/display mode,
    and surrounding context text.
    """
    formulas = []

    for para_idx, para in enumerate(doc.paragraphs):
        elem = para._element

        # Check for display-mode formulas: <m:oMathPara> wrapping
        omath_paras = elem.findall(f".//{_OMATH_PARA}")
        for omp in omath_paras:
            xml_str = etree.tostring(omp, encoding="unicode")
            context = _get_context(doc.paragraphs, para_idx)
            formulas.append(FormulaRef(
                omml_xml=xml_str,
                is_inline=False,
                context_text=context,
            ))

        # Check for inline formulas: bare <m:oMath> not inside <m:oMathPara>
        all_omath = elem.findall(f".//{_OMATH}")
        for om in all_omath:
            # Skip if this oMath is inside an oMathPara (already handled above)
            parent = om.getparent()
            if parent is not None and parent.tag == _OMATH_PARA:
                continue
            xml_str = etree.tostring(om, encoding="unicode")
            context = _get_context(doc.paragraphs, para_idx)
            formulas.append(FormulaRef(
                omml_xml=xml_str,
                is_inline=True,
                context_text=context,
            ))

    log.debug(f"DOCX formula detection: found {len(formulas)} formulas")
    return formulas


def _get_context(paragraphs, para_idx: int, window: int = 2) -> str:
    """Get surrounding paragraph text as context."""
    start = max(0, para_idx - window)
    end = min(len(paragraphs), para_idx + window + 1)
    parts = []
    for i in range(start, end):
        text = paragraphs[i].text.strip()
        if text:
            parts.append(text)
    return " ".join(parts)[:300]
